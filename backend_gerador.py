import os
import subprocess
import unicodedata
from datetime import datetime

import pandas as pd
from docx import Document
from PyPDF2 import PdfReader, PdfWriter


def normalizar(texto):
    if texto is None:
        return ""

    texto = str(texto)

    texto = ''.join(
        c for c in unicodedata.normalize('NFD', texto)
        if unicodedata.category(c) != 'Mn'
    )

    return texto.upper().strip()


def iterar_paragrafos(doc):
    for p in doc.paragraphs:
        yield p

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    yield p


def substituir_textos_doc(doc, substituicoes):
    for p in iterar_paragrafos(doc):
        texto = p.text

        for antigo, novo in substituicoes.items():
            if antigo in texto:
                texto = texto.replace(antigo, novo)

        p.text = texto


def localizar_excel_unico(pasta_entrada):
    if not os.path.exists(pasta_entrada):
        raise FileNotFoundError(f"Pasta de entrada não encontrada: {pasta_entrada}")

    arquivos_xlsx = sorted([
        os.path.join(pasta_entrada, f)
        for f in os.listdir(pasta_entrada)
        if f.lower().endswith(".xlsx")
    ])

    if len(arquivos_xlsx) == 0:
        raise FileNotFoundError(
            f"Nenhum arquivo .xlsx encontrado na pasta de entrada: {pasta_entrada}"
        )

    if len(arquivos_xlsx) > 1:
        nomes = [os.path.basename(x) for x in arquivos_xlsx]
        raise Exception(
            "Foi encontrado mais de um arquivo .xlsx na pasta de entrada.\n"
            f"Arquivos encontrados: {nomes}\n"
            "Deixe apenas um arquivo Excel na pasta."
        )

    return arquivos_xlsx[0]


def localizar_modelos(pasta_modelos):
    if not os.path.exists(pasta_modelos):
        raise FileNotFoundError(f"Pasta de modelos não encontrada: {pasta_modelos}")

    modelo_advertencia = None
    modelo_suspensao = None

    for arquivo in os.listdir(pasta_modelos):
        if not arquivo.lower().endswith(".docx"):
            continue

        nome_norm = normalizar(arquivo)
        caminho = os.path.join(pasta_modelos, arquivo)

        if any(chave in nome_norm for chave in ["ADVERTENCIA", "ADVERT", "ADV"]):
            if modelo_advertencia is None:
                modelo_advertencia = caminho

        if any(chave in nome_norm for chave in ["SUSPENSAO", "SUSP"]):
            if modelo_suspensao is None:
                modelo_suspensao = caminho

    if modelo_advertencia is None:
        raise FileNotFoundError(
            "Nenhum modelo de advertência encontrado na pasta de modelos."
        )

    if modelo_suspensao is None:
        raise FileNotFoundError(
            "Nenhum modelo de suspensão encontrado na pasta de modelos."
        )

    return modelo_advertencia, modelo_suspensao


def converter_pasta_docx_para_pdf(origem, destino, logs):
    arquivos = sorted([
        f for f in os.listdir(origem)
        if f.lower().endswith(".docx")
    ])

    if not arquivos:
        logs.append(f"Nenhum DOCX encontrado em: {origem}")
        return

    for arquivo in arquivos:
        caminho = os.path.join(origem, arquivo)

        comando = [
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            caminho,
            "--outdir", destino
        ]

        resultado = subprocess.run(
            comando,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )

        if resultado.returncode == 0:
            logs.append(f"PDF convertido: {arquivo}")
        else:
            logs.append(f"Erro ao converter {arquivo}: {resultado.stderr}")


def gerar_pdf_final(pasta_pdf, arquivo_saida, copias, logs):
    writer = PdfWriter()

    arquivos_pdf = sorted([
        f for f in os.listdir(pasta_pdf)
        if f.lower().endswith(".pdf")
    ])

    if not arquivos_pdf:
        logs.append(f"Nenhum PDF em {pasta_pdf}")
        return None

    for arquivo in arquivos_pdf:
        caminho = os.path.join(pasta_pdf, arquivo)

        try:
            reader = PdfReader(caminho)

            for pagina in reader.pages:
                for _ in range(copias):
                    writer.add_page(pagina)

        except Exception as erro:
            logs.append(f"Erro ao ler PDF {arquivo}: {erro}")

    with open(arquivo_saida, "wb") as f:
        writer.write(f)

    logs.append(f"PDF final gerado: {arquivo_saida}")
    return arquivo_saida


def processar_documentos(
    pasta_entrada,
    pasta_modelos,
    pasta_saida,
    opcao,
    copias,
    callback_progresso=None
):
    logs = []

    excel_path = localizar_excel_unico(pasta_entrada)
    modelo_advertencia, modelo_suspensao = localizar_modelos(pasta_modelos)

    logs.append(f"Excel localizado: {excel_path}")
    logs.append(f"Modelo de advertência localizado: {modelo_advertencia}")
    logs.append(f"Modelo de suspensão localizado: {modelo_suspensao}")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    raiz_saida = os.path.join(pasta_saida, f"execucao_{timestamp}")

    pasta_susp_docx = os.path.join(raiz_saida, "SUSPENSOES")
    pasta_adv_docx = os.path.join(raiz_saida, "ADVERTENCIAS")
    pasta_susp_pdf = os.path.join(raiz_saida, "PDFS_SUSPENSOES")
    pasta_adv_pdf = os.path.join(raiz_saida, "PDFS_ADVERTENCIAS")

    for pasta in [pasta_susp_docx, pasta_adv_docx, pasta_susp_pdf, pasta_adv_pdf]:
        os.makedirs(pasta, exist_ok=True)

    logs.append(f"Pasta da execução: {raiz_saida}")

    df = pd.read_excel(excel_path)
    
    total_registros = len(df)
    
    if callback_progresso:
        callback_progresso(
            0,
            f"Iniciando processamento de {total_registros} registros..."
        )

    colunas_obrigatorias = ["MATRICULA", "NOME", "UNIDADE", "TURNO", "PUNIÇÃO", "DATA"]
    faltando = [c for c in colunas_obrigatorias if c not in df.columns]

    if faltando:
        raise ValueError(f"Colunas ausentes no Excel: {faltando}")

    meses = {
        "01": "Janeiro",
        "02": "Fevereiro",
        "03": "Março",
        "04": "Abril",
        "05": "Maio",
        "06": "Junho",
        "07": "Julho",
        "08": "Agosto",
        "09": "Setembro",
        "10": "Outubro",
        "11": "Novembro",
        "12": "Dezembro"
    }

    total_advertencias = 0
    total_suspensoes = 0
    sem_data = 0

    for indice, (_, row) in enumerate(df.iterrows(), start=1):
    
        if callback_progresso:
            percentual = int((indice / total_registros) * 100)
    
            callback_progresso(
                percentual,
                f"Processando {indice} de {total_registros}..."
            )
    
        matricula = str(row["MATRICULA"]).strip()
        nome = str(row["NOME"]).strip()
        unidade = str(row["UNIDADE"]).strip()
        turno = str(row["TURNO"]).strip()
        punicao = str(row["PUNIÇÃO"]).strip()
    
        data = pd.to_datetime(row["DATA"], errors="coerce")
        possui_data = not pd.isna(data)
    
        if not possui_data:
            sem_data += 1

        substituicoes_base = {
            "MATRICULA - NOME": f"{matricula} - {nome}",
            "UNIDADE - TURNO": f"{unidade} - {turno}",
        }

        if possui_data:
            dia = data.strftime("%d")
            mes_num = data.strftime("%m")
            ano = data.strftime("%Y")
            mes_extenso = meses[mes_num]

            substituicoes_base["DATA_CAMPINAS"] = f"{dia} de {mes_extenso} de {ano}"
            substituicoes_base["DATA_INICIO"] = data.strftime("%d/%m/%Y")

        nome_limpo = nome.replace(" ", "_")

        if normalizar(punicao) == "ADVERTENCIA ESCRITA":
            if opcao == "suspensoes":
                continue

            doc = Document(modelo_advertencia)
            substituir_textos_doc(doc, substituicoes_base)

            nome_arquivo = os.path.join(
                pasta_adv_docx,
                f"ADVERTENCIA_{matricula}_{nome_limpo}.docx"
            )

            doc.save(nome_arquivo)
            total_advertencias += 1
            logs.append(f"Advertência gerada: {nome_arquivo}")

        else:
            if opcao == "advertencias":
                continue

            dias_punicao = (
                punicao
                .replace("Suspensão", "")
                .replace("Suspensao", "")
                .strip()
            )

            doc = Document(modelo_suspensao)

            substituicoes_susp = substituicoes_base.copy()
            substituicoes_susp["DIAS_PUNICAO"] = dias_punicao

            substituir_textos_doc(doc, substituicoes_susp)

            nome_arquivo = os.path.join(
                pasta_susp_docx,
                f"SUSPENSAO_{matricula}_{nome_limpo}.docx"
            )

            doc.save(nome_arquivo)
            total_suspensoes += 1
            logs.append(f"Suspensão gerada: {nome_arquivo}")

    if opcao in ["suspensoes", "ambos"]:
        converter_pasta_docx_para_pdf(pasta_susp_docx, pasta_susp_pdf, logs)

    if opcao in ["advertencias", "ambos"]:
        converter_pasta_docx_para_pdf(pasta_adv_docx, pasta_adv_pdf, logs)

    data_hoje = datetime.now().strftime("%d-%m-%Y")
    pdf_final_susp = os.path.join(raiz_saida, f"SUSPENSOES_FINAL_{data_hoje}.pdf")
    pdf_final_adv = os.path.join(raiz_saida, f"ADVERTENCIAS_FINAL_{data_hoje}.pdf")

    saidas_pdf = []

    if opcao in ["suspensoes", "ambos"] and total_suspensoes > 0:
        arquivo = gerar_pdf_final(pasta_susp_pdf, pdf_final_susp, copias, logs)
        if arquivo:
            saidas_pdf.append(arquivo)

    if opcao in ["advertencias", "ambos"] and total_advertencias > 0:
        arquivo = gerar_pdf_final(pasta_adv_pdf, pdf_final_adv, copias, logs)
        if arquivo:
            saidas_pdf.append(arquivo)

    saidas_pdf = list(dict.fromkeys(saidas_pdf))

    resumo = {
        "excel_localizado": excel_path,
        "modelo_advertencia_localizado": modelo_advertencia,
        "modelo_suspensao_localizado": modelo_suspensao,
        "pasta_execucao": raiz_saida,
        "pdfs_finais": saidas_pdf,
        "total_advertencias": total_advertencias,
        "total_suspensoes": total_suspensoes,
        "linhas_excel": len(df),
        "registros_sem_data": sem_data,
        "logs": logs
    }

    return resumo
    
